import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os
import math
import tempfile
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont, ImageOps

# --- CONFIGURATION ---
st.set_page_config(
    page_title="Easy Collage", 
    page_icon="✨", 
    layout="wide"
)

# --- STYLING ---
st.markdown("""
<style>
    /* Custom Title */
    .title {
        text-align: center;
        color: #4A90E2;
        font-weight: 800;
        padding: 10px;
    }
    
    /* Make the preview container scrollable on mobile */
    .preview-container {
        height: 400px;
        overflow-y: auto;
        border: 2px dashed #ccc;
        border-radius: 10px;
        padding: 10px;
        background-color: #f9f9f9;
        text-align: center;
        display: flex;
        align-items: center;
        justify-content: center;
    }
    
    /* Large Buttons for Mobile */
    .stButton>button {
        width: 100%;
        height: 3em;
        border-radius: 8px;
        font-weight: bold;
        font-size: 1.1rem;
        margin-top: 5px;
    }
</style>
""", unsafe_allow_html=True)

# --- HELPER FUNCTIONS ---

def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin in ["top", "start", "bottom", "end"]:
        if margin in kwargs:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), str(kwargs[margin]))
            tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_vertical_align(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '12') 
        element.set(qn('w:color'), '000000')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def create_preview(images, title):
    """Creates a low-res PNG for the user to check before downloading."""
    if not images: return None
    
    W, H = 800, 1131 # Low res for speed
    t_h = 100 if title else 0
    avail_h = H - t_h
    canvas = Image.new('RGB', (W, H), 'white')
    
    if title:
        draw = ImageDraw.Draw(canvas)
        try: font = ImageFont.truetype("arial.ttf", 50)
        except: font = ImageFont.load_default()
        bbox = draw.textbbox((0, 0), title, font=font)
        txt_w = bbox[2] - bbox[0]
        txt_h = bbox[3] - bbox[1]
        draw.text(((W-txt_w)/2, (t_h-txt_h)/2), title, font=font, fill="black")

    num = len(images)
    if num == 1: c, r = 1, 1
    elif num <= 4: c, r = 2, 2
    elif num <= 9: c, r = 3, 3
    elif num <= 16: c, r = 4, 4
    else:
        c = math.ceil(math.sqrt(num))
        r = math.ceil(num / c)
        
    cw = W // c
    ch = avail_h // r
    
    idx = 0
    for i in range(r):
        for j in range(c):
            if idx >= num: break
            try:
                img = Image.open(BytesIO(images[idx]))
                img = ImageOps.exif_transpose(img)
                img = img.resize((cw, ch), Image.Resampling.LANCZOS)
                d = ImageDraw.Draw(img)
                d.rectangle([0, 0, cw, ch], outline="black", width=2)
                canvas.paste(img, (j*cw, t_h + i*ch))
            except: pass
            idx += 1
            
    return canvas

# --- FINAL GENERATORS ---

def create_word_doc(images, title):
    doc = Document()
    section = doc.sections[0]
    page_w, page_h = 8.0, 11.0
    num = len(images)
    if num == 0: return None

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.size = Pt(36)
        r.bold = True
        p.space_after = Pt(6)
        h_offset = 0.6
    else:
        h_offset = 0
        
    avail_h = page_h - h_offset

    if num == 1: c, r = 1, 1
    elif num <= 4: c, r = 2, 2
    elif num <= 9: c, r = 3, 3
    elif num <= 16: c, r = 4, 4
    else:
        c = math.ceil(math.sqrt(num))
        r = math.ceil(num / c)
        
    cell_w = (page_w / c)
    cell_h = (avail_h / r)

    table = doc.add_table(rows=r, cols=c)
    table.autofit = False
    for col in table.columns: col.width = Inches(page_w / c)

    idx = 0
    for i in range(r):
        for j in range(c):
            if idx >= num: break
            cell = table.rows[i].cells[j]
            set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
            set_cell_vertical_align(cell, "center")
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            try:
                pil_img = Image.open(BytesIO(images[idx]))
                pil_img = ImageOps.exif_transpose(pil_img)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
                    pil_img.save(tmp)
                    tmp_name = tmp.name
                p.add_run().add_picture(tmp_name, width=Inches(cell_w), height=Inches(cell_h))
                os.remove(tmp_name)
            except: pass
            idx += 1

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def create_png(images, title):
    num = len(images)
    if num == 0: return None
    W, H = 2480, 3508
    t_h = 400 if title else 0
    avail_h = H - t_h
    canvas = Image.new('RGB', (W, H), 'white')
    
    if title:
        draw = ImageDraw.Draw(canvas)
        try: font = ImageFont.truetype("arial.ttf", 160)
        except: font = ImageFont.load_default()
        bbox = draw.textbbox((0, 0), title, font=font)
        txt_w = bbox[2] - bbox[0]
        txt_h = bbox[3] - bbox[1]
        draw.text(((W-txt_w)/2, (t_h-txt_h)/2), title, font=font, fill="black")

    if num == 1: c, r = 1, 1
    elif num <= 4: c, r = 2, 2
    elif num <= 9: c, r = 3, 3
    elif num <= 16: c, r = 4, 4
    else:
        c = math.ceil(math.sqrt(num))
        r = math.ceil(num / c)
        
    cw = W // c
    ch = avail_h // r
    
    idx = 0
    for i in range(r):
        for j in range(c):
            if idx >= num: break
            try:
                img = Image.open(BytesIO(images[idx]))
                img = ImageOps.exif_transpose(img)
                img = img.resize((cw, ch), Image.Resampling.LANCZOS)
                d = ImageDraw.Draw(img)
                d.rectangle([0, 0, cw, ch], outline="black", width=5)
                canvas.paste(img, (j*cw, t_h + i*ch))
            except: pass
            idx += 1
            
    f = BytesIO()
    canvas.save(f, format="PNG")
    f.seek(0)
    return f

# --- MAIN APP ---

st.markdown('<h1 class="title">✨ Easy Photo Collage</h1>', unsafe_allow_html=True)

# Input Section
with st.container():
    col1, col2 = st.columns([2, 1])
    with col1:
        title_input = st.text_input("📝 Enter Title", value="My Collage", label_visibility="collapsed")
    with col2:
        # Just for layout balance or could add advanced options here
        pass

st.markdown("### 1. Choose Photos")
up_col, cam_col = st.columns(2)
with up_col:
    uploaded_files = st.file_uploader("📂 Gallery", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True, label_visibility="collapsed")
with cam_col:
    camera_files = st.camera_input("📸 Camera", label_visibility="collapsed")

all_images_raw = []
if uploaded_files:
    all_images_raw.extend([f.read() for f in uploaded_files])
if camera_files:
    all_images_raw.append(camera_files.read())

# Preview Section
st.markdown("### 2. Preview")
if all_images_raw:
    st.success(f"{len(all_images_raw)} photos loaded.")
    preview_img = create_preview(all_images_raw, title_input)
    if preview_img:
        st.image(preview_img, use_column_width=True)
        
    # Download Section
    st.markdown("### 3. Download")
    col_down1, col_down2 = st.columns(2)
    
    with col_down1:
        if st.button("📄 Download Word", use_container_width=True):
            with st.spinner("Generating..."):
                res = create_word_doc(all_images_raw, title_input)
                if res:
                    st.download_button("⬇️ Save Word", res, f"{title_input}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with col_down2:
        if st.button("🖼️ Download PNG", use_container_width=True):
            with st.spinner("Rendering..."):
                res = create_png(all_images_raw, title_input)
                if res:
                    st.download_button("⬇️ Save PNG", res, f"{title_input}.png", mime="image/png")
else:
    st.info("👇 Upload photos or take a picture to start creating!")

st.markdown("---")
st.markdown("<center><small>Made with ❤️ | Mobile Friendly</small></center>", unsafe_allow_html=True)
